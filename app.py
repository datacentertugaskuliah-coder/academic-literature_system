import streamlit as st
import pandas as pd
import json
import datetime
import re
import os
from io import BytesIO

# ================= DEPENDENCY CHECK =================
def check_dependencies():
    """Cek ketersediaan library opsional tanpa crash"""
    missing = []
    try: import openai
    except ImportError: missing.append("openai (pip install openai)")
    try: import docx
    except ImportError: missing.append("python-docx (pip install python-docx)")
    return missing

DEPS_MISSING = check_dependencies()
if DEPS_MISSING:
    st.warning("⚠️ Library opsional belum terinstall. Fitur AI/Export akan dinonaktifkan sementara.\n" + "\n".join(DEPS_MISSING))

if "openai" not in globals():
    class DummyOpenAI:
        def __init__(self, *args, **kwargs): pass
        class chat:
            class completions:
                @staticmethod
                def create(*args, **kwargs): raise ImportError("openai not installed")
    OpenAI = DummyOpenAI
else:
    from openai import OpenAI

if "docx" in globals():
    import docx
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

# ================= KONFIGURASI HALAMAN =================
st.set_page_config(
    page_title="ALAS v3.0 — Academic Literature System",
    page_icon="📘",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS Mobile-First & Academic Styling
st.markdown("""
<style>
    .stButton button { min-height: 44px; font-weight: 500; }
    .stTabs [data-baseweb="tab-list"] button { padding: 8px 16px; }
    .badge { padding: 4px 8px; border-radius: 6px; font-size: 0.75rem; display: inline-block; }
    .badge-done { background: #dcfce7; color: #16a34a; }
    .badge-pending { background: #fef9c3; color: #ca8a04; }
    .badge-locked { background: #f3f4f6; color: #9ca3af; }
    .guardrail-box { padding: 10px; background: #fffbeb; border-left: 4px solid #f59e0b; border-radius: 4px; margin: 8px 0; font-size: 0.9rem; }
    .academic-text { line-height: 1.8; font-size: 1.05rem; }
    .flowchart-ascii { background: #f1f5f9; padding: 12px; border-radius: 6px; font-family: monospace; font-size: 0.85rem; overflow-x: auto; white-space: pre; }
    .status-ai { background: #dbeafe; border-left: 4px solid #3b82f6; padding: 8px; border-radius: 4px; margin: 8px 0; }
    .status-manual { background: #fffbeb; border-left: 4px solid #f59e0b; padding: 8px; border-radius: 4px; margin: 8px 0; }
</style>
""", unsafe_allow_html=True)

# ================= KONSTANTA & SPESIFIKASI MODUL (Dashboard v2.2) =================
MODULES = ["CL", "M0", "M1", "M2", "M3", "M4", "M5", "M6", "M7", "M8", "M9"]

MODULE_SPECS = {
    "CL": {"name": "Core Layer Fondasi", "desc": "Abstract-Only • Anti-Halusinasi • Evidence Chain • Iteratif"},
    "M0": {"name": "Literature Search", "desc": "Jurnal primer • 4 tahun • 50% terkini • Exclude conference/SLR"},
    "M1": {"name": "Intake Protocol", "desc": "Auto-deteksi metadata • Tanpa PDF • Validasi DOI"},
    "M2": {"name": "Contradiction Finder", "desc": "Kontradiksi genuine antar paper"},
    "M3": {"name": "Citation Chain", "desc": "Genealogi konsep teoritis"},
    "M4": {"name": "Gap Scanner", "desc": "5 research gap ter-ranking"},
    "M5": {"name": "Methodology Audit", "desc": "Evaluasi 4 kriteria metodologi"},
    "M6": {"name": "10 Rekomendasi Judul", "desc": "2500 kata/judul • Latar belakang • Urgensi • GAP • 2 novelty • Flowchart ASCII"},
    "M7": {"name": "Hibah & Publikasi", "desc": "3500 kata • PDP/PFR/Prototype/Model/BRIN • Q1-Q3/SINTA 2-4"},
    "M8": {"name": "Template IMRAD", "desc": "5000 kata • 6 pernyataan/sekssi • Sitasi hanya di Results"},
    "M9": {"name": "Rekomendasi Dataset", "desc": "5 dataset • Colab-ready • Open access • Anti-fake"}
}

RESEARCH_DOMAINS = ["Sosial dan Humaniora", "Sains", "Teknologi", "Ilmu Komputer", "Umum"]
GRANT_SCHEMES = ["PDP", "PFR", "Terapan - Luaran Prototype", "Terapan - Luaran Model", "BRIN"]
PUBLICATION_TARGETS = ["Proposal Penelitian", "Scopus Q1", "Scopus Q2", "Scopus Q3", "SINTA 2", "SINTA 3", "SINTA 4"]

DOMAIN_SPECS = {
    "Sosial dan Humaniora": "Fokus: teori sosial, metode kualitatif/kuantitatif, konteks kebijakan Indonesia",
    "Sains": "Fokus: metode eksperimental, validitas data, reproducibility",
    "Teknologi": "Fokus: implementasi sistem, evaluasi kinerja, studi kasus instansi",
    "Ilmu Komputer": "Fokus: algoritma, kompleksitas, benchmark dataset",
    "Umum": "Fokus: metodologi umum, adaptabilitas lintas disiplin"
}

DEFAULT_CONFIG = {"domain": "Umum", "level": "Skripsi", "target": "Proposal Penelitian", "terminology": {"instansi": True, "pegawai": True}}

# ================= STATE MANAGEMENT (SAFE) =================
def init_state():
    keys_to_check = ["modules", "config", "current", "anchor", "process_mode", "evidence_chain", "m6_batch"]
    for k in keys_to_check:
        if k not in st.session_state:
            st.session_state[k] = None
            
    if st.session_state.modules is None:
        st.session_state.modules = {m: {"status": "locked", "data": {}, "output": None} for m in MODULES}
        st.session_state.modules["M0"]["status"] = "pending"
        st.session_state.modules["CL"]["status"] = "done"
    if st.session_state.config is None:
        st.session_state.config = DEFAULT_CONFIG.copy()
    if st.session_state.current is None:
        st.session_state.current = "M0"
    if st.session_state.anchor is None:
        st.session_state.anchor = {"topic": "", "locked_after_M6": False}
    if st.session_state.process_mode is None:
        st.session_state.process_mode = "🤖 Qwen AI"
    if st.session_state.evidence_chain is None:
        st.session_state.evidence_chain = []
    if st.session_state.m6_batch is None:
        st.session_state.m6_batch = {"start": 1, "size": 1}

def save_state():
    """Update status & unlock next module secara aman"""
    curr = st.session_state.current
    if curr in st.session_state.modules:
        st.session_state.modules[curr]["status"] = "done"
        idx = MODULES.index(curr)
        if idx < len(MODULES) - 1:
            nxt = MODULES[idx + 1]
            if st.session_state.modules[nxt]["status"] == "locked":
                st.session_state.modules[nxt]["status"] = "pending"
    if curr == "M6" and not st.session_state.anchor["locked_after_M6"]:
        st.session_state.anchor["locked_after_M6"] = True
        st.session_state.anchor["topic"] = st.session_state.anchor.get("topic") or "Penelitian berbasis metadata abstrak"

def validate_m0_paper(paper: dict) -> dict:
    """Validasi deterministik M0 dengan safe access"""
    res = {"valid": True, "reasons": [], "flags": []}
    required = ["TITLE", "ABSTRACT", "KEYWORDS", "YEAR", "JOURNAL", "DOI"]
    for f in required:
        if not paper.get(f, "").strip():
            res["valid"] = False; res["reasons"].append(f"Missing: {f}")
    if not res["valid"]: return res
    try:
        y = int(paper.get("YEAR", 0))
        cy = datetime.datetime.now().year
        if not (cy-4 <= y <= cy): res["valid"] = False; res["reasons"].append(f"Year {y} out of range")
        elif y >= cy-1: res["flags"].append("recent")
    except: res["valid"] = False; res["reasons"].append("Invalid YEAR")
    j = paper.get("JOURNAL", "").lower(); a = paper.get("ABSTRACT", "").lower()
    excl = ["conference", "proceedings", "workshop", "systematic review", "meta-analysis", "bibliometric"]
    if any(k in j or k in a for k in excl):
        res["valid"] = False; res["reasons"].append("Excluded: conference/review")
    doi = paper.get("DOI", "")
    if doi and not re.match(r'^10\.\d{4,9}/[-._;()/:A-Z0-9]+$', doi, re.I):
        res["flags"].append("doi_warn")
    return res

# ================= AI & PROMPT ENGINE (DEFENSIVE) =================
def get_qwen_client():
    try:
        key = st.secrets.get("QWEN_API_KEY", "")
        base = st.secrets.get("QWEN_BASE_URL", "https://dashscope.aliyuncs.com/compatible-mode/v1")
        if not key: return None
        return OpenAI(api_key=key, base_url=base)
    except Exception: return None

def generate_prompt_for_module(mod_id: str, inp: dict) -> str:
    cfg = st.session_state.config
    dom_note = DOMAIN_SPECS.get(cfg.get("domain", "Umum"), "")
    term_note = "\n• Gunakan: 'instansi' (bukan 'perusahaan'), 'pegawai' (bukan 'karyawan')" if cfg.get("terminology", {}).get("instansi") else ""
    spec = MODULE_SPECS.get(mod_id, {}).get("desc", "")
    
    base = f"""[CORE_LAYER_v3.0 — {'QWEN' if st.session_state.process_mode=='🤖 Qwen AI' else 'MANUAL'} MODE]
Bidang: {cfg['domain']} | Level: {cfg['level']} | Target: {cfg['target']}{term_note}
Spesifikasi: {spec} | Konteks: {dom_note}
[INPUT]\n{json.dumps(inp, indent=2, ensure_ascii=False)}
[INSTRUKSI]
1. Analisis HANYA dari metadata (judul, abstrak, keyword). Dilarang akses PDF/full-text.
2. Jangan mengarang data/sitasi/metode. Gunakan tag: `[Ref: ID_Paper, Field: Abstract/Keyword]`
3. Jika ambigu: `[NEEDS_CLARIFICATION: <field>]`. Jangan asumsi.
4. Output JSON STRICT: {{"layer1_summary":"...", "layer2_academic":"...", "layer3_metadata":{{"module":"{mod_id}", "guardrails":true}}}}
[STATUS_MODUL: {mod_id}_SELESAI]"""
    return base

def run_ai_processing(mod_id: str, inp: dict) -> dict:
    """Semi-auto dengan fallback manual & error handling"""
    client = get_qwen_client()
    prompt = generate_prompt_for_module(mod_id, inp)
    
    if client and st.session_state.process_mode == "🤖 Qwen AI" and "openai" in globals():
        try:
            with st.spinner("🤖 Qwen memproses... (10-30dtk)"):
                res = client.chat.completions.create(
                    model=st.secrets.get("QWEN_MODEL", "qwen-plus"),
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.1, max_tokens=4000,
                    response_format={"type": "json_object"}
                )
                raw = res.choices[0].message.content.strip()
                if raw.startswith("```json"): raw = raw.split("```json")[1].split("```")[0]
                elif raw.startswith("```"): raw = raw.split("```")[1].split("```")[0]
                out = json.loads(raw)
                out["status"] = "ai_success"; out["mode"] = "qwen"
                return out
        except json.JSONDecodeError:
            st.warning("⚠️ Respons JSON tidak valid. Fallback manual.")
        except Exception as e:
            st.warning(f"⚠️ API Error: {str(e)[:100]}. Fallback manual.")

    # Fallback Manual
    if mod_id == "M0" and isinstance(inp.get("raw"), str):
        blocks = [b.strip() for b in inp["raw"].split('\n\n') if b.strip()]
        valid_papers, recent = [], 0
        for i, blk in enumerate(blocks[:10]):
            p = {k.strip().upper(): v.strip() for k,v in (line.split(':',1) for line in blk.split('\n') if ':' in line)}
            val = validate_m0_paper(p)
            p["_id"] = f"P{i+1}"
            if val["valid"]:
                valid_papers.append(p)
                if "recent" in val["flags"]: recent += 1
        inp["validated_papers"] = valid_papers
        inp["summary"] = f"{len(valid_papers)} valid. {recent}/{len(valid_papers)} terkini."
        if valid_papers and recent/len(valid_papers) < 0.5:
            st.warning("⚠️ <50% paper dari 1-2 tahun terkini.")

    return {
        "status": "manual_mode", "mode": "fallback",
        "layer1_summary": "⚠️ Mode Manual: Salin prompt di tab 📖 Detail, paste ke AI, kembalikan hasil.",
        "layer2_academic": prompt,
        "layer3_metadata": {"module": mod_id, "mode": "manual", "guardrails": ["Anti-Halusinasi", "Evidence Chain", "Abstract-Only"]}
    }

# ================= EXPORT ENGINE (SAFE) =================
def generate_academic_docx(target: str, title: str, content: str, refs: list) -> BytesIO:
    if "docx" not in globals():
        raise ImportError("python-docx tidak terinstall. Export dinonaktifkan.")
    doc = docx.Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)
    
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Target: {target} | Bidang: {st.session_state.config['domain']} | Tanggal: {datetime.date.today()}").italic = True
    doc.add_page_break()
    
    doc.add_heading("Isi Dokumen", level=2)
    doc.add_paragraph(content if content else "Konten akan dilengkapi setelah validasi manual.")
    
    doc.add_heading("Evidence Chain & Disclaimer", level=2)
    doc.add_paragraph("Catatan: Seluruh klaim bersumber dari metadata abstrak intake. Verifikasi manual wajib dilakukan sebelum submit/ajuan.")
    for r in refs[-10:]:
        doc.add_paragraph(f"[{r.get('ref','')}] → {r.get('module','')} | {r.get('timestamp','')[:10]}", style='List Bullet')
        
    buf = BytesIO(); doc.save(buf); buf.seek(0); return buf

# ================= SIDEBAR =================
with st.sidebar:
    st.header("📘 ALAS v3.0")
    st.session_state.process_mode = st.radio("Mode Inferensi", ["🤖 Qwen AI (Semi-Otomatis)", "👐 Manual (Eksternal AI)"], index=0)
    if st.session_state.process_mode == "🤖 Qwen AI":
        st.success("✅ Terhubung Qwen API") if get_qwen_client() else st.error("❌ QWEN_API_KEY tidak ditemukan di secrets")
    st.divider()
    
    st.subheader("🔧 Konfigurasi")
    with st.expander("📝 Terminologi"):
        st.session_state.config["terminology"]["instansi"] = st.checkbox("Gunakan 'instansi'", value=st.session_state.config["terminology"]["instansi"])
        st.session_state.config["terminology"]["pegawai"] = st.checkbox("Gunakan 'pegawai'", value=st.session_state.config["terminology"]["pegawai"])
    st.session_state.config["domain"] = st.selectbox("Bidang", RESEARCH_DOMAINS, index=RESEARCH_DOMAINS.index(st.session_state.config["domain"]) if st.session_state.config["domain"] in RESEARCH_DOMAINS else 0)
    st.session_state.config["level"] = st.selectbox("Level", ["Skripsi", "Tesis", "Disertasi", "Hibah"], index=0)
    st.session_state.config["target"] = st.selectbox("Target Publikasi", PUBLICATION_TARGETS, index=PUBLICATION_TARGETS.index(st.session_state.config["target"]) if st.session_state.config["target"] in PUBLICATION_TARGETS else 0)
    st.divider()
    
    done = sum(1 for m in st.session_state.modules.values() if m["status"]=="done")
    st.progress(done/len(MODULES)); st.caption(f"{done}/{len(MODULES)} selesai")
    st.session_state.current = st.radio("Pilih Modul", MODULES, 
        format_func=lambda x: f"{x} • {'✅' if st.session_state.modules[x]['status']=='done' else '🔒' if st.session_state.modules[x]['status']=='locked' else '⏳'}")
    st.divider()
    
    if st.button("💾 Backup State"):
        st.download_button("Download JSON", json.dumps(st.session_state.to_dict(), indent=2), f"ALAS_backup_{datetime.date.today()}.json")
    if st.button("🗑️ Reset", type="primary"):
        for m in MODULES: st.session_state.modules[m] = {"status": "locked", "data": {}, "output": None}
        st.session_state.modules["M0"]["status"] = "pending"; st.session_state.modules["CL"]["status"] = "done"
        st.session_state.current = "M0"; st.session_state.anchor = {"topic": "", "locked_after_M6": False}
        st.session_state.evidence_chain = []
        st.rerun()

# ================= RENDER FUNCTIONS =================
def render_output_layer(mod_id):
    out = st.session_state.modules[mod_id].get("output")
    if not out: st.warning("⏳ Belum diproses."); return
    t1, t2, t3 = st.tabs(["📋 Ringkasan", "📖 Detail", "📦 Metadata/Export"])
    with t1: st.markdown(out.get("layer1_summary", ""))
    with t2:
        if out.get("status") == "ai_success":
            st.markdown('<div class="status-ai">🤖 Hasil Qwen AI — review & edit sebelum simpan</div>', unsafe_allow_html=True)
            txt = st.text_area("Edit Hasil", value=out.get("layer2_academic",""), height=500 if mod_id in ["M6","M7","M8"] else 400, key=f"edit_{mod_id}")
            if st.button("💾 Simpan & Validasi", type="primary"):
                if txt.strip():
                    st.session_state.modules[mod_id]["output"]["layer2_academic"] = txt
                    st.session_state.modules[mod_id]["output"]["user_verified"] = True
                    refs = re.findall(r'\[Ref:\s*([^\]]+)\]', txt)
                    for r in refs: st.session_state.evidence_chain.append({"module": mod_id, "ref": r, "timestamp": datetime.datetime.now().isoformat()})
                    save_state(); st.success("✅ Tersimpan."); st.rerun()
                else: st.warning("⚠️ Hasil tidak boleh kosong.")
            if st.button("🔄 Regenerasi"): st.session_state.modules[mod_id]["output"] = None; st.rerun()
        else:
            st.markdown('<div class="status-manual">👐 Mode Manual: Salin prompt → AI → Paste hasil → Simpan</div>', unsafe_allow_html=True)
            st.code(out.get("layer2_academic", ""), language="text")
            txt = st.text_area("Paste Hasil AI", value="", height=500, key=f"manual_{mod_id}")
            if st.button("💾 Simpan", type="primary"):
                if txt.strip():
                    st.session_state.modules[mod_id]["output"]["layer2_academic"] = txt
                    st.session_state.modules[mod_id]["output"]["user_verified"] = True
                    for r in re.findall(r'\[Ref:\s*([^\]]+)\]', txt): st.session_state.evidence_chain.append({"module": mod_id, "ref": r, "timestamp": datetime.datetime.now().isoformat()})
                    save_state(); st.success("✅ Tersimpan."); st.rerun()
    with t3:
        st.json(out.get("layer3_metadata", {}))
        if "docx" in globals() and mod_id in ["M7","M8"] and st.session_state.modules[mod_id].get("output", {}).get("user_verified"):
            st.markdown("### 📥 Export .docx (Rekomendasi High-Impact)")
            top3 = ["Proposal Penelitian", "Scopus Q2", "SINTA 2"]  # Simple rule-based top3
            cols = st.columns(3)
            for i, t in enumerate(top3):
                with cols[i]:
                    st.info(f"🎯 {t}")
                    if st.button(f"📄 Download {t}"):
                        try:
                            buf = generate_academic_docx(t, st.session_state.anchor.get("topic","Penelitian"), out.get("layer2_academic","")[:2000], st.session_state.evidence_chain)
                            st.download_button(f"✅ {t}.docx", buf, f"ALAS_{t.replace(' ','_')}_{datetime.date.today()}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                        except ImportError: st.error("python-docx belum terinstall.")

def render_generic(mod_id):
    st.header(f"⚙️ {mod_id}: {MODULE_SPECS[mod_id]['name']}")
    st.caption(MODULE_SPECS[mod_id]['desc'])
    if st.session_state.modules[mod_id]["status"] == "locked":
        st.error("🔒 Selesaikan modul sebelumnya."); return
    st.markdown('<div class="guardrail-box">🛡️ Guardrail: Abstract-Only • Anti-Halusinasi • Evidence Chain</div>', unsafe_allow_html=True)
    if not st.session_state.modules[mod_id].get("output"):
        if st.button(f"🚀 Proses {mod_id}", type="primary"):
            inp = st.session_state.modules[mod_id].get("data", {})
            if mod_id in ["M7","M8","M9"] and st.session_state.anchor.get("locked_after_M6"): inp["_anchor"] = st.session_state.anchor
            st.session_state.modules[mod_id]["output"] = run_ai_processing(mod_id, inp)
            save_state(); st.success("✅ Siap review."); st.rerun()
    else:
        st.success(f"✅ {mod_id} selesai ({st.session_state.modules[mod_id]['output'].get('mode','manual')})")
    render_output_layer(mod_id)

def render_m0():
    st.header("📥 M0: Literature Search")
    with st.form("m0"):
        kw = st.text_input("Keyword", placeholder="Contoh: sentiment analysis")
        txt = st.text_area("Metadata Paper (format: TITLE: ...)", height=250)
        if st.form_submit_button("✅ Validasi & Proses", type="primary"):
            if not kw or not txt: st.error("❌ Wajib diisi."); return
            st.session_state.modules["M0"]["data"] = {"keyword": kw, "raw": txt}
            st.session_state.modules["M0"]["output"] = run_ai_processing("M0", st.session_state.modules["M0"]["data"])
            save_state(); st.success("✅ Prompt/Validasi siap."); st.rerun()
    if st.session_state.modules["M0"].get("output"): render_output_layer("M0")

# ================= MAIN ROUTER =================
init_state()
mod = st.session_state.current
if mod == "CL":
    st.header("🔷 Core Layer — Aktif ✅")
    for m in MODULES:
        st.markdown(f"- {'✅' if st.session_state.modules[m]['status']=='done' else '🔒' if st.session_state.modules[m]['status']=='locked' else '⏳'} **{m}**: {MODULE_SPECS[m]['name']}")
elif mod == "M0": render_m0()
elif mod in ["M1","M2","M3","M4","M5"]:
    st.info(f"🚧 {mod} dalam pengembangan framework. Gunakan M0, M6-M9 untuk alur lengkap."); render_generic(mod)
elif mod == "M6":
    st.header("🎯 M6: 10 Rekomendasi Judul"); st.caption("2500 kata/judul • Latar belakang • Urgensi • GAP • 2 novelty • Flowchart")
    if st.session_state.modules["M5"]["status"] != "done": st.error("🔒 M5 wajib selesai."); st.stop()
    with st.form("m6cfg"):
        st.session_state.m6_batch["size"] = st.slider("Judul per batch", 1, 3, 1)
        st.form_submit_button("🚀 Generate", type="primary")
        inp = {"m5_audit": st.session_state.modules["M5"].get("data",{}), "batch": st.session_state.m6_batch}
        st.session_state.modules["M6"]["output"] = run_ai_processing("M6", inp)
        save_state(); st.success("✅ Siap review."); st.rerun()
    if st.session_state.modules["M6"].get("output"): render_output_layer("M6")
elif mod in ["M7","M8","M9"]: render_generic(mod)

st.divider()
st.caption(f"📘 ALAS v3.0 • Abstract-Only • Dashboard v2.2 • Qwen Semi-Auto • Export Ready • Bidang: {st.session_state.config['domain']}")
